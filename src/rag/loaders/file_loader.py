"""文件加载器 - 支持PDF/MD/TXT/DOCX"""
import logging
from pathlib import Path

from rag.loaders.base import BaseLoader, Document

logger = logging.getLogger(__name__)


class FileLoader(BaseLoader):
    """文件加载器"""

    SUPPORTED_EXTENSIONS = {".pdf", ".md", ".txt", ".docx"}

    def __init__(self, encoding: str = "utf-8"):
        self.encoding = encoding

    def load(self, source: str) -> list[Document]:
        """加载文件或目录"""
        path = Path(source)

        if not path.exists():
            raise FileNotFoundError(f"路径不存在: {source}")

        if path.is_file():
            return [self._load_file(path)]
        elif path.is_dir():
            return self._load_directory(path)
        else:
            raise ValueError(f"不支持的路径类型: {source}")

    def _load_file(self, file_path: Path) -> Document:
        """加载单个文件"""
        ext = file_path.suffix.lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件类型: {ext}")

        if ext == ".pdf":
            content = self._load_pdf(file_path)
        elif ext == ".docx":
            content = self._load_docx(file_path)
        else:
            content = self._load_text(file_path)

        metadata = self._get_file_metadata(file_path)
        metadata["file_type"] = ext[1:]  # 去掉点号

        return Document(
            content=content,
            metadata=metadata,
            source=str(file_path),
        )

    def _load_directory(self, dir_path: Path) -> list[Document]:
        """加载目录下所有支持的文件"""
        documents = []

        for ext in self.SUPPORTED_EXTENSIONS:
            for file_path in dir_path.rglob(f"*{ext}"):
                try:
                    doc = self._load_file(file_path)
                    documents.append(doc)
                except Exception as e:
                    logger.warning(f"加载文件失败 {file_path}: {e}")

        return documents

    def _load_text(self, file_path: Path) -> str:
        """加载纯文本文件"""
        with open(file_path, "r", encoding=self.encoding) as f:
            return f.read()

    def _load_pdf(self, file_path: Path) -> str:
        """加载PDF文件"""
        from PyPDF2 import PdfReader

        reader = PdfReader(str(file_path))
        text_parts = []

        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)

        return "\n\n".join(text_parts)

    def _load_docx(self, file_path: Path) -> str:
        """加载DOCX文件"""
        from docx import Document as DocxDocument

        doc = DocxDocument(str(file_path))
        paragraphs = [para.text for para in doc.paragraphs if para.text]
        return "\n\n".join(paragraphs)
